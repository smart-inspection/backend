from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from sqlalchemy.orm import Session

from app.services.report_template_service import build_company_report_context

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUTPUT_DIR = Path("output/reports")

def _safe_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default

def _dict_value(data: dict[str, Any], *keys: str, default: Any = None) -> Any:
    for key in keys:
        if key in data and data[key] is not None:
            value = data[key]
            if isinstance(value, str):
                value = value.strip()
                if value == "":
                    continue
            return value
    return default

def _escape_pdf_text(value: Any, default: str = "--") -> str:
    return (
        _safe_text(value, default)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )

def _normalize_filename(value: str) -> str:
    text = _safe_text(value, "reporte").lower().strip()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = re.sub(r"-{2,}", "-", text).strip("-")
    return text or "reporte"

def _ensure_parent(path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    return path

def _existing_image(path_value: str | None) -> str | None:
    if not path_value:
        return None
    path = Path(path_value)
    if path.exists() and path.is_file():
        return str(path)
    return None

def _chunked(items: list[Any], size: int) -> list[list[Any]]:
    return [items[i:i + size] for i in range(0, len(items), size)]

def _slot_status_text(item: dict[str, Any]) -> str:
    return "PRESENTE" if bool(_dict_value(item, "present", default=False)) else "PENDIENTE"

def _slot_status_color_hex(item: dict[str, Any]) -> str:
    return "1B9E5A" if bool(_dict_value(item, "present", default=False)) else "C62828"

def _slot_status_fill_hex(item: dict[str, Any]) -> str:
    return "EAF6EA" if bool(_dict_value(item, "present", default=False)) else "FDECEC"

# =========================
# DOCX helpers
# =========================

def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(_safe_text(text, "--"))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def _shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue

        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)

        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))

def _apply_table_borders(table, color: str = "000000", size: int = 8):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top={"val": "single", "sz": size, "color": color},
                bottom={"val": "single", "sz": size, "color": color},
                left={"val": "single", "sz": size, "color": color},
                right={"val": "single", "sz": size, "color": color},
            )

def _add_page_number_fields(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(8)

    fld_page_begin = OxmlElement("w:fldChar")
    fld_page_begin.set(qn("w:fldCharType"), "begin")

    instr_page = OxmlElement("w:instrText")
    instr_page.set(qn("xml:space"), "preserve")
    instr_page.text = " PAGE "

    fld_page_end = OxmlElement("w:fldChar")
    fld_page_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_page_begin)
    run._r.append(instr_page)
    run._r.append(fld_page_end)

    run2 = paragraph.add_run(" de ")
    run2.font.name = "Arial"
    run2.font.size = Pt(8)

    run3 = paragraph.add_run()
    run3.font.name = "Arial"
    run3.font.size = Pt(8)

    fld_total_begin = OxmlElement("w:fldChar")
    fld_total_begin.set(qn("w:fldCharType"), "begin")

    instr_total = OxmlElement("w:instrText")
    instr_total.set(qn("xml:space"), "preserve")
    instr_total.text = " NUMPAGES "

    fld_total_end = OxmlElement("w:fldChar")
    fld_total_end.set(qn("w:fldCharType"), "end")

    run3._r.append(fld_total_begin)
    run3._r.append(instr_total)
    run3._r.append(fld_total_end)

def _configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Arial"

def _add_paragraph(
    document_or_cell,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    size: int = 10,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    color: str | None = None,
    space_after: int = 4,
    space_before: int = 0,
):
    p = document_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(_safe_text(text, ""))
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def _add_divider(document):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def _add_footer_docx(document: Document, context: dict[str, Any]):
    footer_data = context["footer"]
    section = document.sections[0]
    footer = section.footer

    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(footer_data["company_line"])
    r1.font.name = "Arial"
    r1.font.size = Pt(8)

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(footer_data["address_line"])
    r2.font.name = "Arial"
    r2.font.size = Pt(8)

    p3 = footer.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(footer_data["contact_line"])
    r3.font.name = "Arial"
    r3.font.size = Pt(8)

    p4 = footer.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(footer_data["website_line"])
    r4.font.name = "Arial"
    r4.font.size = Pt(8)
    r4.underline = True

    p5 = footer.add_paragraph()
    _add_page_number_fields(p5)

def _add_cover_docx(document: Document, context: dict[str, Any]):
    header = context["header"]
    branding = context["branding"]

    logo_path = _existing_image(header.get("logo_path"))
    if logo_path:
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(logo_path, width=Inches(2.2))

    _add_paragraph(
        document,
        branding["report_title"],
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
        space_before=2,
    )
    _add_paragraph(
        document,
        branding["report_code_display"],
        bold=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=2,
    )
    _add_paragraph(
        document,
        branding["report_subtitle"],
        bold=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    _add_paragraph(
        document,
        branding["equipment_display"],
        bold=True,
        underline=True,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )

    _add_paragraph(document, f"TIPO DE INSPECCIÓN: {header['inspection_type']}", bold=True, size=10)
    _add_paragraph(document, f"FECHA DE INSPECCIÓN: {header['inspection_date']}", bold=True, size=10)
    _add_paragraph(document, f"MÉTODOS EMPLEADOS: {header['methods_display']}", bold=True, size=10)
    _add_paragraph(document, f"CONDICIÓN GENERAL DEL EQUIPO: {header['general_condition']}", bold=True, size=10)

    _add_paragraph(
        document,
        _safe_text(header.get("location"), "").upper(),
        bold=True,
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
        space_before=10,
    )

    document.add_page_break()

def _add_technical_intro_docx(document: Document, context: dict[str, Any]):
    tech = context["technical_info"]

    _add_paragraph(
        document,
        "INFORME TÉCNICO",
        bold=True,
        underline=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run("SOLICITADO POR: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r2 = p.add_run(tech["requested_by"])
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run("DIRECCIÓN: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r2 = p.add_run(tech["address"])
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run("RESPONSABLE DEL SERVICIO: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r2 = p.add_run(tech["service_responsible"])
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r1 = p.add_run("FECHA DE INSPECCIÓN: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r2 = p.add_run(tech["inspection_date_long"])
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

    _add_divider(document)
    _add_paragraph(document, tech["intro_paragraph"], size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)

def _add_heading_docx(document: Document, title: str):
    _add_paragraph(
        document,
        title,
        bold=True,
        size=11.5,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=6,
        space_after=6,
    )

def _add_identification_docx(document: Document, context: dict[str, Any]):
    data = context["identification"]

    _add_heading_docx(document, "1. IDENTIFICACIÓN DEL EQUIPO INSPECCIONADO")

    lines = [
        ("Tipo de Equipo", data["tipo_equipo"]),
        ("N° de placa", data["placa"]),
        ("Marca", data["marca"]),
        ("N° de VIN", data["vin"]),
        ("Año de fabricación", data["anio_fabricacion"]),
        ("Kilometraje de Referencia", data["kilometraje"]),
        ("Antigüedad", data["antiguedad"]),
        ("N° de Ejes", data["numero_ejes"]),
        ("Carga Útil", data["carga_util"]),
        ("Peso Neto", data["peso_neto"]),
        ("Marca de King pin", data["marca_king_pin"]),
        ("Modelo de King Pin", data["modelo_king_pin"]),
        ("N° de Serie de King pin", data["serie_king_pin"]),
    ]

    for label, value in lines:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10.5)
        r2 = p.add_run(_safe_text(value, "--"))
        r2.font.name = "Arial"
        r2.font.size = Pt(10.5)

def _add_bullets_docx(document: Document, items: list[str]):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(_safe_text(item, "--"))
        run.font.name = "Arial"
        run.font.size = Pt(10.5)

def _add_frequency_table_docx(document: Document, context: dict[str, Any]):
    _add_heading_docx(document, "5. FRECUENCIA DE INSPECCIÓN")

    rows = context["frequency_table"]
    equipment_type = _safe_text(context["identification"]["tipo_equipo"], "EQUIPO").upper()

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = [equipment_type, "Método de Inspección", "Frecuencia (*)", "Porcentaje"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(cell, "EDEDED")

    for row in rows:
        r = table.add_row().cells
        _set_cell_text(r[0], row["component"], size=10)
        _set_cell_text(r[1], row["method"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(r[2], row["frequency"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(r[3], row["percentage"], size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    _apply_table_borders(table)
    _add_paragraph(document, context["frequency"], size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6)

def _add_equipment_table_docx(document: Document, context: dict[str, Any]):
    _add_heading_docx(document, "7. EQUIPOS DE INSPECCIÓN EMPLEADOS")

    equipment = context["inspection_equipment"]
    mt_items = equipment["mt"]
    vt_items = equipment["vt"]
    max_rows = max(len(mt_items), len(vt_items)) + 1

    table = document.add_table(rows=max_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    _set_cell_text(table.rows[0].cells[0], equipment["mt_title"], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(table.rows[0].cells[1], equipment["vt_title"], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(table.rows[0].cells[0], "EDEDED")
    _shade_cell(table.rows[0].cells[1], "EDEDED")

    for i in range(max_rows - 1):
        _set_cell_text(table.rows[i + 1].cells[0], mt_items[i] if i < len(mt_items) else "", size=10)
        _set_cell_text(table.rows[i + 1].cells[1], vt_items[i] if i < len(vt_items) else "", size=10)

    _apply_table_borders(table)

def _add_criteria_docx(document: Document, context: dict[str, Any]):
    _add_heading_docx(document, "8. CRITERIOS DE INSPECCIÓN")
    criteria = context["criteria"]

    p1 = document.add_paragraph()
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run("ACEPTADO: ")
    r1.bold = True
    r1.italic = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x1B, 0x9E, 0x5A)
    r2 = p1.add_run(criteria["accepted"])
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

    p2 = document.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    r1 = p2.add_run("RECHAZADO: ")
    r1.bold = True
    r1.italic = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)
    r2 = p2.add_run(criteria["rejected"])
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

    for item in criteria.get("rejected_items", []):
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)

    p3 = document.add_paragraph()
    p3.paragraph_format.space_before = Pt(6)
    r1 = p3.add_run("RETIRO DE LA OPERACIÓN: ")
    r1.bold = True
    r1.italic = True
    r1.font.name = "Arial"
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = RGBColor(0xD4, 0xA3, 0x00)
    r2 = p3.add_run(criteria["retirement"])
    r2.italic = True
    r2.font.name = "Arial"
    r2.font.size = Pt(10.5)

def _add_results_table_docx(document: Document, context: dict[str, Any]):
    _add_heading_docx(document, "9. RESULTADOS DE LA INSPECCIÓN")

    rows = context["results"]
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Equipo", "Componente", "Condición", "Observaciones", "Acción Correctiva"]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, text, bold=True, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(cell, "EDEDED")

    for row in rows:
        r = table.add_row().cells
        _set_cell_text(r[0], row["equipo"], size=9.5)
        _set_cell_text(r[1], row["componente"], size=9.5)
        _set_cell_text(r[2], row["condicion"], size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(r[3], row["observaciones"], size=9.5)
        _set_cell_text(r[4], row["accion"], size=9.5)

    _apply_table_borders(table)

# =========================
# DOCX evidence rendering
# =========================

def _render_docx_slot_image(cell, image_path: str | None, slot_label: str, present: bool):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if image_path:
        run = p.add_run()
        try:
            run.add_picture(image_path, width=Inches(2.85))
            return
        except Exception:
            pass

    label = "EVIDENCIA PENDIENTE" if not present else "IMAGEN NO DISPONIBLE"
    run = p.add_run(label)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(_safe_text(slot_label, "Evidencia esperada"))
    run2.font.name = "Arial"
    run2.font.size = Pt(8.5)

def _add_fixed_evidence_sections_docx_document(document: Document, context: dict[str, Any]) -> bool:
    fixed_sections = list(context.get("fixed_evidence_sections") or [])
    if not fixed_sections:
        return False

    document.add_page_break()
    _add_paragraph(
        document,
        "RESULTADOS",
        bold=True,
        underline=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    for section in fixed_sections:
        section_title = _safe_text(
            _dict_value(section, "section_title", "sectiontitle"),
            "Sección de evidencias",
        )

        _add_paragraph(
            document,
            section_title,
            bold=True,
            size=11,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=4,
            space_after=6,
        )

        items = list(_dict_value(section, "items", default=[]) or [])
        if not items:
            _add_paragraph(
                document,
                "No se definieron evidencias esperadas para esta sección.",
                size=9.5,
                align=WD_ALIGN_PARAGRAPH.LEFT,
                space_after=4,
            )
            continue

        for item in items:
            present = bool(_dict_value(item, "present", default=False))
            slot_label = _safe_text(_dict_value(item, "slot_label", "slotlabel"), "Evidencia esperada")
            caption = _safe_text(_dict_value(item, "caption"), "No registrada")
            file_type = _safe_text(_dict_value(item, "file_type", "filetype"), "No registrado")
            evidence_category = _safe_text(_dict_value(item, "evidence_category", "evidencecategory"), "No registrado")
            ocr_text = _safe_text(_dict_value(item, "ocr_text", "ocrtext"), "Sin OCR")
            file_path = _dict_value(item, "file_path", "filepath")
            image_path = _existing_image(file_path)

            table = document.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            _render_docx_slot_image(left_cell, image_path, slot_label, present)

            _shade_cell(right_cell, _slot_status_fill_hex(item))
            _add_paragraph(right_cell, slot_label, bold=True, size=10, space_after=2)
            _add_paragraph(
                right_cell,
                f"Estado: {_slot_status_text(item)}",
                bold=True,
                size=9,
                color=_slot_status_color_hex(item),
                space_after=2,
            )
            _add_paragraph(right_cell, f"Descripción: {caption}", size=9, space_after=1)
            _add_paragraph(right_cell, f"Categoría: {evidence_category}", size=9, space_after=1)
            _add_paragraph(right_cell, f"Tipo de archivo: {file_type}", size=9, space_after=1)
            _add_paragraph(right_cell, f"Ruta: {_safe_text(file_path, 'Sin ruta')}", size=8.5, space_after=1)
            _add_paragraph(right_cell, f"OCR: {ocr_text}", size=8.5, space_after=1)

            _apply_table_borders(table)
            document.add_paragraph()

        document.add_paragraph()

    return True

def _add_evidences_docx_fallback(document: Document, context: dict[str, Any]):
    evidences = context.get("evidences") or []
    if not evidences:
        return

    document.add_page_break()
    _add_paragraph(
        document,
        "RESULTADOS",
        bold=True,
        underline=True,
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )

    for chunk in _chunked(evidences, 2):
        table = document.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for col_index in range(2):
            if col_index >= len(chunk):
                _set_cell_text(table.rows[0].cells[col_index], "", size=9)
                _set_cell_text(table.rows[1].cells[col_index], "", size=9)
                continue

            evidence = chunk[col_index]
            image_path = _existing_image(evidence.get("path"))
            top_cell = table.rows[0].cells[col_index]
            bottom_cell = table.rows[1].cells[col_index]

            top_cell.text = ""
            p = top_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if image_path:
                run = p.add_run()
                try:
                    run.add_picture(image_path, width=Inches(2.75))
                except Exception:
                    p.add_run("Imagen no disponible")
            else:
                p.add_run("Imagen no disponible")

            caption = evidence.get("display_title") or evidence.get("caption") or f"Foto {evidence.get('index', '')}"
            _set_cell_text(bottom_cell, caption, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        _apply_table_borders(table)
        document.add_paragraph()

def _add_evidences_docx_document(document: Document, context: dict[str, Any]):
    if _add_fixed_evidence_sections_docx_document(document, context):
        return
    _add_evidences_docx_fallback(document, context)

def _build_docx_document(context: dict[str, Any]) -> Document:
    document = Document()
    _configure_document(document)
    _add_footer_docx(document, context)

    signature_path = _existing_image(context["technical_info"].get("signature_path"))

    _add_cover_docx(document, context)
    _add_technical_intro_docx(document, context)

    _add_identification_docx(document, context)

    _add_heading_docx(document, "2. OBJETIVO")
    _add_paragraph(document, context["objective"], size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    _add_heading_docx(document, "3. ALCANCE")
    _add_bullets_docx(document, context["scope"])

    _add_heading_docx(document, "4. PROTOCOLO EMPLEADO")
    _add_paragraph(document, context["protocol"], size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    _add_frequency_table_docx(document, context)

    _add_heading_docx(document, "6. NORMAS Y CÓDIGOS DE REFERENCIA")
    _add_bullets_docx(document, context["standards"])

    _add_equipment_table_docx(document, context)
    _add_criteria_docx(document, context)
    _add_results_table_docx(document, context)

    _add_heading_docx(document, "10. CONCLUSIONES")
    _add_paragraph(document, context["conclusion"], size=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    if signature_path:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(signature_path, width=Inches(1.8))

        _add_paragraph(
            document,
            context["technical_info"]["service_responsible"],
            bold=True,
            size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=1,
        )
        _add_paragraph(
            document,
            "RESPONSABLE DEL SERVICIO",
            size=9,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=8,
        )

    _add_evidences_docx_document(document, context)
    return document

# =========================
# PDF helpers
# =========================

def _pdf_styles():
    styles = getSampleStyleSheet()

    styles.add(
        ParagraphStyle(
            name="PdfTitleCenter",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=14,
            leading=17,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PdfCodeCenter",
            parent=styles["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=3,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PdfBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            alignment=TA_JUSTIFY,
            textColor=colors.black,
            spaceAfter=5,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PdfBodyLeft",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            alignment=TA_LEFT,
            textColor=colors.black,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PdfSection",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=13,
            alignment=TA_LEFT,
            textColor=colors.black,
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="PdfSmallCenter",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=3,
        )
    )
    return styles

def _rl_paragraph(text: str, style, *, allow_markup: bool = False):
    text = _safe_text(text, "--")
    if not allow_markup:
        text = _escape_pdf_text(text, "--")
    return Paragraph(text, style)

def _rl_table(
    data,
    widths,
    header_fill="#EDEDED",
    body_fill_1="#FFFFFF",
    body_fill_2="#F7F7F7",
    fontsize=8.6,
):
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header_fill)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), fontsize),
                ("LEADING", (0, 0), (-1, -1), fontsize + 2),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor(body_fill_1), colors.HexColor(body_fill_2)]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table

def _pdf_footer(canvas, doc, context: dict[str, Any]):
    footer = context["footer"]
    width, _ = A4

    canvas.saveState()
    canvas.setFont("Helvetica", 7.5)
    canvas.drawCentredString(width / 2, 1.6 * cm, footer["company_line"])
    canvas.drawCentredString(width / 2, 1.25 * cm, footer["address_line"])
    canvas.drawCentredString(width / 2, 0.92 * cm, footer["contact_line"])
    canvas.drawCentredString(width / 2, 0.59 * cm, footer["website_line"])
    canvas.drawRightString(width - 1.6 * cm, 0.32 * cm, f"{canvas.getPageNumber()}")
    canvas.restoreState()

# =========================
# PDF evidence rendering
# =========================

def _build_fixed_evidence_sections_pdf(context: dict[str, Any], styles) -> list[Any]:
    fixed_sections = list(context.get("fixed_evidence_sections") or [])
    if not fixed_sections:
        return []

    story: list[Any] = [
        _rl_paragraph("<u><b>RESULTADOS</b></u>", styles["PdfTitleCenter"], allow_markup=True),
        Spacer(1, 0.15 * cm),
    ]

    for section in fixed_sections:
        section_title = _safe_text(_dict_value(section, "section_title", "sectiontitle"), "Sección de evidencias")
        story.append(_rl_paragraph(section_title, styles["PdfSection"]))

        items = list(_dict_value(section, "items", default=[]) or [])
        if not items:
            story.append(_rl_paragraph("No se definieron evidencias esperadas para esta sección.", styles["PdfBody"]))
            story.append(Spacer(1, 0.12 * cm))
            continue

        for item in items:
            present = bool(_dict_value(item, "present", default=False))
            slot_label = _safe_text(_dict_value(item, "slot_label", "slotlabel"), "Evidencia esperada")
            caption = _safe_text(_dict_value(item, "caption"), "No registrada")
            file_type = _safe_text(_dict_value(item, "file_type", "filetype"), "No registrado")
            evidence_category = _safe_text(_dict_value(item, "evidence_category", "evidencecategory"), "No registrado")
            ocr_text = _safe_text(_dict_value(item, "ocr_text", "ocrtext"), "Sin OCR")
            file_path = _dict_value(item, "file_path", "filepath")
            image_path = _existing_image(file_path)

            if image_path:
                left_cell = RLImage(image_path, width=7.0 * cm, height=5.0 * cm)
            else:
                missing_label = "<b>IMAGEN NO DISPONIBLE</b>" if present else "<b>EVIDENCIA PENDIENTE</b>"
                left_cell = _rl_paragraph(missing_label, styles["PdfSmallCenter"], allow_markup=True)

            status_text = "PRESENTE" if present else "PENDIENTE"
            status_color = "1B9E5A" if present else "C62828"

            right_cell = [
                _rl_paragraph(
                    f"<b>Elemento esperado:</b> {_escape_pdf_text(slot_label)}",
                    styles["PdfBodyLeft"],
                    allow_markup=True,
                ),
                _rl_paragraph(
                    f"<b>Estado:</b> <font color='#{status_color}'>{status_text}</font>",
                    styles["PdfBodyLeft"],
                    allow_markup=True,
                ),
                _rl_paragraph(
                    f"<b>Descripción:</b> {_escape_pdf_text(caption)}",
                    styles["PdfBodyLeft"],
                    allow_markup=True,
                ),
                _rl_paragraph(
                    f"<b>Categoría:</b> {_escape_pdf_text(evidence_category)}",
                    styles["PdfBodyLeft"],
                    allow_markup=True,
                ),
                _rl_paragraph(
                    f"<b>Tipo:</b> {_escape_pdf_text(file_type)}",
                    styles["PdfBodyLeft"],
                    allow_markup=True,
                ),
            ]

            slot_table = Table([[left_cell, right_cell]], colWidths=[7.2 * cm, 9.4 * cm])
            slot_table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#EAF6EA" if present else "#FDECEC")),
                    ]
                )
            )

            story.append(KeepTogether([slot_table, Spacer(1, 0.18 * cm)]))

        story.append(Spacer(1, 0.15 * cm))

    return story

def _build_pdf_evidences_fallback(context: dict[str, Any], styles) -> list[Any]:
    evidences = context.get("evidences") or []
    if not evidences:
        return []

    story: list[Any] = [
        _rl_paragraph("<u><b>RESULTADOS</b></u>", styles["PdfTitleCenter"], allow_markup=True),
        Spacer(1, 0.15 * cm),
    ]

    for chunk in _chunked(evidences, 2):
        image_row = []
        caption_row = []

        for evidence in chunk:
            img_path = _existing_image(evidence.get("path"))
            if img_path:
                image_row.append(RLImage(img_path, width=7.3 * cm, height=5.2 * cm))
            else:
                image_row.append(_rl_paragraph("Imagen no disponible", styles["PdfSmallCenter"]))

            caption = evidence.get("display_title") or evidence.get("caption") or f"Foto {evidence.get('index')}"
            caption_row.append(
                _rl_paragraph(
                    f"<b>{_escape_pdf_text(caption)}</b>",
                    styles["PdfSmallCenter"],
                    allow_markup=True,
                )
            )

        if len(image_row) == 1:
            image_row.append("")
            caption_row.append("")

        table = Table([image_row, caption_row], colWidths=[8.0 * cm, 8.0 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.append(KeepTogether([table, Spacer(1, 0.2 * cm)]))

    return story

def _build_pdf_story(context: dict[str, Any]):
    styles = _pdf_styles()
    story = []

    header = context["header"]
    branding = context["branding"]
    tech = context["technical_info"]
    identification = context["identification"]

    signature_path = _existing_image(context["technical_info"].get("signature_path"))

    logo_path = _existing_image(header.get("logo_path"))
    if logo_path:
        story.append(RLImage(logo_path, width=4.8 * cm, height=2.1 * cm))
        story.append(Spacer(1, 0.15 * cm))

    story.append(_rl_paragraph(_safe_text(branding["report_title"]), styles["PdfTitleCenter"]))
    story.append(_rl_paragraph(_safe_text(branding["report_code_display"]), styles["PdfCodeCenter"]))
    story.append(_rl_paragraph(_safe_text(branding["report_subtitle"]), styles["PdfCodeCenter"]))
    story.append(Spacer(1, 0.15 * cm))
    story.append(
        _rl_paragraph(
            f"<u>{_escape_pdf_text(branding['equipment_display'])}</u>",
            styles["PdfCodeCenter"],
            allow_markup=True,
        )
    )

    story.append(
        _rl_paragraph(
            f"<b>TIPO DE INSPECCIÓN:</b> {_escape_pdf_text(header['inspection_type'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<b>FECHA DE INSPECCIÓN:</b> {_escape_pdf_text(header['inspection_date'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<b>MÉTODOS EMPLEADOS:</b> {_escape_pdf_text(header['methods_display'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<b>CONDICIÓN GENERAL DEL EQUIPO:</b> {_escape_pdf_text(header['general_condition'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(Spacer(1, 0.35 * cm))
    story.append(_rl_paragraph(_safe_text(header["location"]), styles["PdfSmallCenter"]))
    story.append(PageBreak())

    story.append(_rl_paragraph("<u><b>INFORME TÉCNICO</b></u>", styles["PdfTitleCenter"], allow_markup=True))
    story.append(
        _rl_paragraph(
            f"<b>SOLICITADO POR:</b> {_escape_pdf_text(tech['requested_by'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<b>DIRECCIÓN:</b> {_escape_pdf_text(tech['address'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<b>RESPONSABLE DEL SERVICIO:</b> {_escape_pdf_text(tech['service_responsible'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<b>FECHA DE INSPECCIÓN:</b> {_escape_pdf_text(tech['inspection_date_long'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(Spacer(1, 0.1 * cm))
    story.append(HRFlowable(width="100%", thickness=0.7, color=colors.black))
    story.append(Spacer(1, 0.15 * cm))
    story.append(_rl_paragraph(_safe_text(tech["intro_paragraph"]), styles["PdfBody"]))

    story.append(_rl_paragraph("1. IDENTIFICACIÓN DEL EQUIPO INSPECCIONADO", styles["PdfSection"]))
    ident_lines = [
        f"<b>Tipo de Equipo:</b> {_escape_pdf_text(identification['tipo_equipo'])}",
        f"<b>N° de placa:</b> {_escape_pdf_text(identification['placa'])}",
        f"<b>Marca:</b> {_escape_pdf_text(identification['marca'])}",
        f"<b>N° de VIN:</b> {_escape_pdf_text(identification['vin'])}",
        f"<b>Año de fabricación:</b> {_escape_pdf_text(identification['anio_fabricacion'])}",
        f"<b>Kilometraje de Referencia:</b> {_escape_pdf_text(identification['kilometraje'])}",
        f"<b>Antigüedad:</b> {_escape_pdf_text(identification['antiguedad'])}",
        f"<b>N° de Ejes:</b> {_escape_pdf_text(identification['numero_ejes'])}",
        f"<b>Carga Útil:</b> {_escape_pdf_text(identification['carga_util'])}",
        f"<b>Peso Neto:</b> {_escape_pdf_text(identification['peso_neto'])}",
        f"<b>Marca de King pin:</b> {_escape_pdf_text(identification['marca_king_pin'])}",
        f"<b>Modelo de King Pin:</b> {_escape_pdf_text(identification['modelo_king_pin'])}",
        f"<b>N° de Serie de King pin:</b> {_escape_pdf_text(identification['serie_king_pin'])}",
    ]
    for line in ident_lines:
        story.append(_rl_paragraph(line, styles["PdfBodyLeft"], allow_markup=True))

    story.append(_rl_paragraph("2. OBJETIVO", styles["PdfSection"]))
    story.append(_rl_paragraph(_safe_text(context["objective"]), styles["PdfBody"]))

    story.append(_rl_paragraph("3. ALCANCE", styles["PdfSection"]))
    for item in context["scope"]:
        story.append(_rl_paragraph(f"- {_safe_text(item)}", styles["PdfBodyLeft"]))

    story.append(_rl_paragraph("4. PROTOCOLO EMPLEADO", styles["PdfSection"]))
    story.append(_rl_paragraph(_safe_text(context["protocol"]), styles["PdfBody"]))

    story.append(_rl_paragraph("5. FRECUENCIA DE INSPECCIÓN", styles["PdfSection"]))
    freq_data = [[
        _safe_text(context["identification"]["tipo_equipo"], "EQUIPO").upper(),
        "Método de Inspección",
        "Frecuencia (*)",
        "Porcentaje",
    ]]
    for row in context["frequency_table"]:
        freq_data.append([row["component"], row["method"], row["frequency"], row["percentage"]])

    story.append(_rl_table(freq_data, [5.8 * cm, 4.4 * cm, 3.4 * cm, 2.8 * cm], fontsize=8.3))
    story.append(Spacer(1, 0.12 * cm))
    story.append(_rl_paragraph(_safe_text(context["frequency"]), styles["PdfBody"]))

    story.append(_rl_paragraph("6. NORMAS Y CÓDIGOS DE REFERENCIA", styles["PdfSection"]))
    for item in context["standards"]:
        story.append(_rl_paragraph(f"• {_safe_text(item)}", styles["PdfBodyLeft"]))

    story.append(_rl_paragraph("7. EQUIPOS DE INSPECCIÓN EMPLEADOS", styles["PdfSection"]))
    equipment = context["inspection_equipment"]
    max_rows = max(len(equipment["mt"]), len(equipment["vt"]))
    eq_data = [[equipment["mt_title"], equipment["vt_title"]]]
    for i in range(max_rows):
        eq_data.append([
            equipment["mt"][i] if i < len(equipment["mt"]) else "",
            equipment["vt"][i] if i < len(equipment["vt"]) else "",
        ])
    story.append(_rl_table(eq_data, [8.0 * cm, 8.0 * cm], fontsize=8.5))

    story.append(_rl_paragraph("8. CRITERIOS DE INSPECCIÓN", styles["PdfSection"]))
    criteria = context["criteria"]
    story.append(
        _rl_paragraph(
            f"<font color='#1B9E5A'><b><i>ACEPTADO:</i></b></font> {_escape_pdf_text(criteria['accepted'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    story.append(
        _rl_paragraph(
            f"<font color='#1F77B4'><b><i>RECHAZADO:</i></b></font> {_escape_pdf_text(criteria['rejected'])}",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )
    for item in criteria.get("rejected_items", []):
        story.append(
            _rl_paragraph(
                f"• <i>{_escape_pdf_text(item)}</i>",
                styles["PdfBodyLeft"],
                allow_markup=True,
            )
        )
    story.append(
        _rl_paragraph(
            f"<font color='#D4A300'><b><i>RETIRO DE LA OPERACIÓN:</i></b></font> <i>{_escape_pdf_text(criteria['retirement'])}</i>",
            styles["PdfBodyLeft"],
            allow_markup=True,
        )
    )

    story.append(_rl_paragraph("9. RESULTADOS DE LA INSPECCIÓN", styles["PdfSection"]))
    results_data = [["Equipo", "Componente", "Condición", "Observaciones", "Acción Correctiva"]]
    for row in context["results"]:
        results_data.append([
            row["equipo"],
            row["componente"],
            row["condicion"],
            row["observaciones"],
            row["accion"],
        ])
    story.append(_rl_table(results_data, [3.0 * cm, 3.8 * cm, 2.2 * cm, 4.0 * cm, 3.0 * cm], fontsize=7.8))

    story.append(_rl_paragraph("10. CONCLUSIONES", styles["PdfSection"]))
    story.append(_rl_paragraph(_safe_text(context["conclusion"]), styles["PdfBody"]))

    if signature_path:
        story.append(Spacer(1, 0.4 * cm))
        story.append(RLImage(signature_path, width=4.0 * cm, height=2.0 * cm))
        story.append(
            _rl_paragraph(f"<b>{tech['service_responsible']}</b>", styles["PdfSmallCenter"], allow_markup=True))
        story.append(_rl_paragraph("RESPONSABLE DEL SERVICIO", styles["PdfSmallCenter"]))

    fixed_evidence_story = _build_fixed_evidence_sections_pdf(context, styles)
    if fixed_evidence_story:
        story.append(PageBreak())
        story.extend(fixed_evidence_story)
    else:
        fallback_story = _build_pdf_evidences_fallback(context, styles)
        if fallback_story:
            story.append(PageBreak())
            story.extend(fallback_story)

    return story

# =========================
# Exports
# =========================

def export_report_docx(db: Session, draft_id: int, output_path: str | Path | None = None) -> str:
    context = build_company_report_context(db, draft_id)

    if output_path is None:
        report_code = _normalize_filename(context["header"]["report_code_display"])
        output_path = OUTPUT_DIR / f"{report_code}.docx"

    output_path = _ensure_parent(Path(output_path))
    document = _build_docx_document(context)
    document.save(str(output_path))
    return str(output_path)

def export_report_pdf(db: Session, draft_id: int, output_path: str | Path | None = None) -> str:
    context = build_company_report_context(db, draft_id)

    if output_path is None:
        report_code = _normalize_filename(context["header"]["report_code_display"])
        output_path = OUTPUT_DIR / f"{report_code}.pdf"

    output_path = _ensure_parent(Path(output_path))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.5 * cm,
        bottomMargin=2.0 * cm,
        title=context["header"]["report_title"],
        author=context["company"]["name"],
    )

    story = _build_pdf_story(context)
    doc.build(
        story,
        onFirstPage=lambda canvas, d: _pdf_footer(canvas, d, context),
        onLaterPages=lambda canvas, d: _pdf_footer(canvas, d, context),
    )
    return str(output_path)

def export_report_files(db: Session, draft_id: int, output_dir: str | Path | None = None) -> dict[str, str]:
    if output_dir is None:
        output_dir = OUTPUT_DIR

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    context = build_company_report_context(db, draft_id)
    report_code = _normalize_filename(context["header"]["report_code_display"])

    docx_path = output_dir / f"{report_code}.docx"
    pdf_path = output_dir / f"{report_code}.pdf"

    export_report_docx(db, draft_id, docx_path)
    export_report_pdf(db, draft_id, pdf_path)

    return {
        "docx": str(docx_path),
        "pdf": str(pdf_path),
    }